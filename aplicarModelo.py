import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Directory containing your Word documents
directory = r'Z:\CARTORIO DE SIMAO DIAS\LIVRO 2-AD - REGISTRO DE IMÓVEIS\Words'

# Path to your model (template) document
model_path = r'Z:\CARTORIO DE SIMAO DIAS\LIVRO 2-AD - REGISTRO DE IMÓVEIS\Words\model.docx'

# Open the model document
model_doc = Document(model_path)

# Function to copy the formatting from the model document
def apply_model_formatting(doc, model_doc):
    # Copy formatting from the model document to the new one
    for para_model in model_doc.paragraphs:
        para = doc.add_paragraph()  # Add a new paragraph to the document

        # Copy the text and formatting (font family, font size, and justification)
        for run_model in para_model.runs:
            run = para.add_run(run_model.text)  # Copy the text
            run.font.name = run_model.font.name  # Copy font family
            run.font.size = run_model.font.size  # Copy font size
        
        # Copy the alignment from the model
        para.alignment = para_model.alignment

# Loop through all Word documents in the specified directory
for filename in os.listdir(directory):
    if filename.endswith('.docx'):
        # Full path of the document
        doc_path = os.path.join(directory, filename)
        
        # Open the current document (this could be empty)
        doc = Document(doc_path)
        
        # Apply the model's formatting
        apply_model_formatting(doc, model_doc)
        
        # Save the modified document
        doc.save(doc_path)

print("Model document formatting applied to all documents.")
